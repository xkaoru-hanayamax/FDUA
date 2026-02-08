"""
DOCX出力モジュール

提案書をDOCX形式で出力
"""

import re
from pathlib import Path
from typing import Optional, Union

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..common.config import Config, default_config
from ..common.debug import debug_log, debug_docx_processing


class DocxWriter:
    """DOCX出力クラス"""

    def __init__(self, config: Optional[Config] = None):
        """
        Args:
            config: 設定オブジェクト
        """
        self.config = config or default_config

    def compose_proposal_text(
        self,
        company_code: str,
        company_info: dict,
        sections: dict[str, str],
    ) -> str:
        """
        全セクションを結合して提案書テキストを構成

        Args:
            company_code: 企業コード
            company_info: 企業基本情報
            sections: セクション辞書

        Returns:
            提案書全体のテキスト
        """
        proposal = f"""成長戦略提案書

対象企業コード: {company_code}
対象企業所在地: {company_info.get('location', '')}
対象企業業種: {company_info.get('industry', '')}

{'=' * 60}

1. 企業概要・分析
{'=' * 60}

{sections.get('overview', '（未生成）')}

{'=' * 60}

2. 課題の抽出
{'=' * 60}

{sections.get('issues', '（未生成）')}

{'=' * 60}

3. 成長戦略・提案
{'=' * 60}

{sections.get('strategy', '（未生成）')}

{'=' * 60}

4. 効果試算
{'=' * 60}

{sections.get('effects', '（未生成）')}

{'=' * 60}

5. ロードマップ
{'=' * 60}

{sections.get('roadmap', '（未生成）')}

"""
        return proposal

    def save_docx(
        self,
        company_code: str,
        company_info: dict,
        sections: dict[str, str],
        output_path: Optional[Union[str, Path]] = None,
    ) -> str:
        """
        提案書をDOCX形式で保存

        Args:
            company_code: 企業コード
            company_info: 企業基本情報
            sections: セクション辞書
            output_path: 出力ファイルパス（Noneの場合はデフォルト）

        Returns:
            保存したファイルパス
        """
        if output_path is None:
            output_path = self.config.get_proposal_path(company_code)
        else:
            output_path = Path(output_path)

        # 出力ディレクトリを作成
        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()

        # タイトル
        title = doc.add_heading("成長戦略提案書", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 基本情報
        doc.add_paragraph(f"対象企業コード: {company_code}")
        doc.add_paragraph(f"対象企業所在地: {company_info.get('location', '')}")
        doc.add_paragraph(f"対象企業業種: {company_info.get('industry', '')}")
        doc.add_paragraph("")

        # 各セクション
        section_list = [
            ("1. 企業概要・分析", sections.get("overview", "")),
            ("2. 課題の抽出", sections.get("issues", "")),
            ("3. 成長戦略・提案", sections.get("strategy", "")),
            ("4. 効果試算", sections.get("effects", "")),
            ("5. ロードマップ", sections.get("roadmap", "")),
        ]

        for heading, content in section_list:
            doc.add_heading(heading, 1)
            debug_log("docx_writer", f"セクション '{heading}' を出力中...", content[:500] if content else "（空）")

            # マークダウンコンテンツをWord要素に変換して追加
            self._parse_and_add_content(doc, content)

        doc.save(str(output_path))
        print(f"提案書を保存しました: {output_path}")

        debug_log("docx_writer", f"DOCX出力完了: {output_path}")

        return str(output_path)

    def save_prompt_log(
        self,
        company_code: str,
        prompt_logs: list[dict],
        output_path: Optional[Union[str, Path]] = None,
    ) -> str:
        """
        プロンプトログを保存

        Args:
            company_code: 企業コード
            prompt_logs: プロンプトログのリスト
            output_path: 出力ファイルパス

        Returns:
            保存したファイルパス
        """
        if output_path is None:
            output_path = self.config.proposals_dir / f"{company_code}_proposal_log.txt"
        else:
            output_path = Path(output_path)

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(f"提案書生成プロンプトログ - 企業コード: {company_code}\n")
            f.write("=" * 80 + "\n\n")

            for i, log in enumerate(prompt_logs, 1):
                f.write(f"[{i}] セクション: {log['section']}\n")
                f.write("-" * 60 + "\n")
                f.write("【入力プロンプト】\n")
                f.write(log["prompt"] + "\n\n")
                f.write("【LLM出力】\n")
                f.write(log["response"] + "\n")
                f.write("\n" + "=" * 60 + "\n\n")

        return str(output_path)

    def _is_table_line(self, line: str) -> bool:
        """表の行かどうかを判定"""
        stripped = line.strip()
        return stripped.startswith("|") and stripped.endswith("|")

    def _collect_table_lines(self, lines: list[str], start: int) -> tuple[list[str], int]:
        """連続する表の行を収集"""
        table_lines = []
        i = start
        while i < len(lines) and self._is_table_line(lines[i]):
            table_lines.append(lines[i])
            i += 1
        return table_lines, i

    def _add_table(self, doc: Document, table_lines: list[str]):
        """マークダウン表をWord表に変換"""
        rows = []
        for line in table_lines:
            # 区切り行（|---|---|）をスキップ
            if re.match(r"^\|[-:\s|]+\|$", line.strip()):
                continue
            # セルを抽出
            cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
            rows.append(cells)

        if not rows:
            return

        # 表を作成
        num_cols = len(rows[0])
        table = doc.add_table(rows=len(rows), cols=num_cols)
        table.style = "Table Grid"

        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < num_cols:
                    cell = table.rows[row_idx].cells[col_idx]
                    cell.text = cell_text
                    # セル内のフォントサイズを設定
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10.5)

    def _add_list_item(self, doc: Document, line: str):
        """箇条書き項目を追加"""
        # インデントレベルを計算
        indent_level = (len(line) - len(line.lstrip())) // 2
        text = line.strip().lstrip("-*").strip()

        p = doc.add_paragraph(text, style="List Bullet")
        # インデント設定
        if indent_level > 0:
            p.paragraph_format.left_indent = Pt(18 * indent_level)
        # フォントサイズ設定
        for run in p.runs:
            run.font.size = Pt(10.5)

    def _parse_and_add_content(self, doc: Document, content: str):
        """マークダウンコンテンツをWord要素に変換して追加"""
        lines = content.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]

            # 表の検出（|で始まり|で終わる行が連続）
            if self._is_table_line(line):
                table_lines, i = self._collect_table_lines(lines, i)
                self._add_table(doc, table_lines)
                continue

            # 見出しの検出（レベル順に判定）
            if line.startswith("### "):
                doc.add_heading(line[4:].strip(), 3)
            elif line.startswith("## "):
                doc.add_heading(line[3:].strip(), 2)
            elif line.startswith("# "):
                # セクション内の#見出しは2に落とす（セクション見出し自体が1なので）
                doc.add_heading(line[2:].strip(), 2)

            # 箇条書きの検出
            elif line.strip().startswith("- ") or line.strip().startswith("* "):
                self._add_list_item(doc, line)

            # 通常段落
            elif line.strip():
                p = doc.add_paragraph(line.strip())
                for run in p.runs:
                    run.font.size = Pt(10.5)

            i += 1

    def count_characters(
        self,
        company_code: str,
        company_info: dict,
        sections: dict[str, str],
    ) -> int:
        """
        提案書の文字数をカウント

        Args:
            company_code: 企業コード
            company_info: 企業基本情報
            sections: セクション辞書

        Returns:
            文字数
        """
        proposal = self.compose_proposal_text(company_code, company_info, sections)
        return len(proposal)
